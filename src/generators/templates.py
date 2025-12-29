"""Шаблоны для различных типов вопросов"""

from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from typing import List, Dict, Optional, Any

from ..models.question import Question
from ..models.metadata import DocumentMetadata


class QuestionTemplate:
    """Базовый класс для шаблонов вопросов"""
    
    @staticmethod
    def set_table_borders(table):
        """
        Установка границ для таблицы
        
        Args:
            table: Объект таблицы Word
        """
        table.style = 'Table Grid'
        
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        tblBorders = OxmlElement('w:tblBorders')
        
        borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        for border_name in borders:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(tblBorders)
    
    @staticmethod
    def add_answer_table(doc: Document, task_number: int, answer: str):
        """
        Добавление таблицы для ответа с границами
        
        Args:
            doc: Объект документа Word
            task_number: Номер задания
            answer: Текст ответа
        """
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        
        # Настройка ширины колонок
        # widths = [Inches(2.0), Inches(0.5), Inches(4.0)]
        # for i, width in enumerate(widths):
        #     table.columns[i].width = width
        
        # Установка границ таблицы
        QuestionTemplate.set_table_borders(table)
        
        # Заполнение таблицы
        cells = table.rows[0].cells
        
        # Колонка 1: Номер вопроса
        cell1 = cells[0]
        cell1.text = f"№{task_number}"
        for paragraph in cell1.paragraphs:
            paragraph.style = doc.styles['AnswerText']
        
        # Колонка 2: Пустая
        cell2 = cells[1]
        for paragraph in cell2.paragraphs:
            paragraph.style = doc.styles['AnswerText']
        
        # Колонка 3: Ответ
        cell3 = cells[2]
        cell3.text = answer
        for paragraph in cell3.paragraphs:
            paragraph.style = doc.styles['AnswerText']

    @staticmethod
    def apply_table_widths(doc: Document, table, widths_pct: List[int]) -> None:
        """
        Применяет ширины колонок таблицы в процентах от доступной ширины страницы.
        """
        if not widths_pct:
            return
        if len(widths_pct) != len(table.columns):
            return
        total = sum(widths_pct)
        if total <= 0:
            return
        section = doc.sections[0]
        available = section.page_width - section.left_margin - section.right_margin
        # normalize
        for i, pct in enumerate(widths_pct):
            table.columns[i].width = available * (float(pct) / float(total))


class EssayGigachatTemplate(QuestionTemplate):
    """Шаблон для вопросов типа essay_gigachat"""
    
    @staticmethod
    def render(doc: Document, question: Question, metadata: DocumentMetadata, task_number: int, cfg: Optional[dict] = None):
        """
        Рендеринг вопроса типа essay_gigachat
        
        Args:
            doc: Объект документа Word
            question: Объект вопроса
            metadata: Метаданные документа
            task_number: Номер задания
        """
        # Заголовок задания
        header_text = (
            f"- Задание {task_number} "
            f"({metadata.pk_prefix}-{metadata.pk_id} – "
            f"{metadata.ipk_prefix}-{metadata.ipk_id} "
            f"{metadata.description})"
        )
        
        doc.add_paragraph(header_text, style='QuestionHeader')
        
        # Текст вопроса
        doc.add_paragraph(question.question_text, style='Question')
        
        # Таблица для ответа
        QuestionTemplate.add_answer_table(doc, task_number, question.reference_answer)
        try:
            cols = ((cfg or {}).get("layout") or {}).get("essay_gigachat", {}).get("table_cols_pct")
            if isinstance(cols, list):
                QuestionTemplate.apply_table_widths(doc, doc.tables[-1], [int(x) for x in cols])
        except Exception:
            pass


class MultichoiceTemplate(QuestionTemplate):
    """Шаблон для вопросов типа multichoice"""
    
    @staticmethod
    def add_multichoice_answer_table(doc: Document, task_number: int, answers: List[str], correct_answers: List[str], cfg: Optional[dict] = None):
        """
        Добавление таблицы для ответов multichoice с границами
        
        Args:
            doc: Объект документа Word
            task_number: Номер задания
            answers: Список всех вариантов ответов
            correct_answers: Список правильных ответов
        """
        # Количество строк: 1 заголовочная + количество вариантов ответов
        num_rows = 1 + len(answers) if answers else 1
        table = doc.add_table(rows=num_rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        
        # Настройка ширины колонок
        # widths = [Inches(2.0), Inches(2.5), Inches(3.5)]
        # for i, width in enumerate(widths):
        #     table.columns[i].width = width
        
        # Установка границ таблицы
        QuestionTemplate.set_table_borders(table)
        
        # Заполнение таблицы
        # Ячейка (0,0): номер вопроса
        cell_0_0 = table.rows[0].cells[0]
        cell_0_0.text = f"№{task_number}"
        for paragraph in cell_0_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
        
        # Ячейка (1,0): "Варианты ответа:" (жирный)
        cell_1_0 = table.rows[0].cells[1]
        cell_1_0.text = "Варианты ответа:"
        for paragraph in cell_1_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
            for run in paragraph.runs:
                run.bold = True
        
        # Ячейка (2,0): "Правильный ответ:" (жирный)
        cell_2_0 = table.rows[0].cells[2]
        cell_2_0.text = "Правильный ответ:"
        for paragraph in cell_2_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
            for run in paragraph.runs:
                run.bold = True
        
        # Заполнение вариантов ответов и правильных ответов
        for i, answer in enumerate(answers, start=1):
            # Ячейка (1,i): вариант ответа
            cell_1_i = table.rows[i].cells[1]
            cell_1_i.text = answer
            for paragraph in cell_1_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']
            
            # Ячейка (2,i): правильный ответ (если этот вариант правильный)
            cell_2_i = table.rows[i].cells[2]
            if answer in correct_answers:
                cell_2_i.text = answer
            else:
                cell_2_i.text = ""  # Пустая строка для неправильных ответов
            for paragraph in cell_2_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']
        
        # Ячейка (0,i): пустая для всех строк с ответами
        for i in range(1, num_rows):
            cell_0_i = table.rows[i].cells[0]
            for paragraph in cell_0_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']

        try:
            cols = ((cfg or {}).get("layout") or {}).get("multichoice", {}).get("table_cols_pct")
            if isinstance(cols, list):
                QuestionTemplate.apply_table_widths(doc, table, [int(x) for x in cols])
        except Exception:
            pass
    
    @staticmethod
    def render(doc: Document, question: Question, metadata: DocumentMetadata, task_number: int, cfg: Optional[dict] = None):
        """
        Рендеринг вопроса типа multichoice
        
        Args:
            doc: Объект документа Word
            question: Объект вопроса
            metadata: Метаданные документа
            task_number: Номер задания
        """
        # Заголовок задания
        header_text = (
            f"- Задание {task_number} "
            f"({metadata.pk_prefix}-{metadata.pk_id} – "
            f"{metadata.ipk_prefix}-{metadata.ipk_id} "
            f"{metadata.description})"
        )
        
        doc.add_paragraph(header_text, style='QuestionHeader')
        
        # Текст вопроса
        doc.add_paragraph(question.question_text, style='Question')
        
        # Таблица для ответов multichoice
        MultichoiceTemplate.add_multichoice_answer_table(
            doc, 
            task_number, 
            question.answers, 
            question.correct_answers,
            cfg,
        )


class MatchingTemplate(QuestionTemplate):
    """Шаблон для вопросов типа matching"""
    
    @staticmethod
    def add_matching_answer_table(doc: Document, task_number: int, matching_items: List[Dict[str, str]], matching_answers: List[str], cfg: Optional[dict] = None):
        """
        Добавление таблицы для ответов matching с границами
        
        Args:
            doc: Объект документа Word
            task_number: Номер задания
            matching_items: Список элементов сопоставления с ответами
            matching_answers: Список всех уникальных вариантов ответов
        """
        # Количество строк: 1 заголовочная + максимальное из (количество элементов сопоставления, количество вариантов ответов)
        max_data_rows = max(len(matching_items), len(matching_answers)) if matching_items or matching_answers else 0
        num_rows = 1 + max_data_rows if max_data_rows > 0 else 1
        
        table = doc.add_table(rows=num_rows, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        
        # Установка границ таблицы
        QuestionTemplate.set_table_borders(table)
        
        # Заполнение заголовочной строки
        # Ячейка (0,0): номер вопроса
        cell_0_0 = table.rows[0].cells[0]
        cell_0_0.text = f"№{task_number}"
        for paragraph in cell_0_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
        
        # Ячейка (1,0): "Варианты ответа:" (жирный)
        cell_1_0 = table.rows[0].cells[1]
        cell_1_0.text = "Варианты ответа:"
        for paragraph in cell_1_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
            for run in paragraph.runs:
                run.bold = True
        
        # Ячейка (2,0): "Элемент для сопоставления:" (жирный)
        cell_2_0 = table.rows[0].cells[2]
        cell_2_0.text = "Элемент для сопоставления:"
        for paragraph in cell_2_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
            for run in paragraph.runs:
                run.bold = True
        
        # Ячейка (3,0): "Правильный ответ:" (жирный)
        cell_3_0 = table.rows[0].cells[3]
        cell_3_0.text = "Правильный ответ:"
        for paragraph in cell_3_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
            for run in paragraph.runs:
                run.bold = True
        
        # Заполнение данных (начиная со строки 1)
        for i in range(1, num_rows):
            row_idx = i - 1  # Индекс для списков (0-based)
            
            # Ячейка (0,i): пустая
            cell_0_i = table.rows[i].cells[0]
            for paragraph in cell_0_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']
            
            # Ячейка (1,i): вариант ответа (если есть)
            if row_idx < len(matching_answers):
                cell_1_i = table.rows[i].cells[1]
                cell_1_i.text = matching_answers[row_idx]
                for paragraph in cell_1_i.paragraphs:
                    paragraph.style = doc.styles['AnswerText']
            else:
                cell_1_i = table.rows[i].cells[1]
                for paragraph in cell_1_i.paragraphs:
                    paragraph.style = doc.styles['AnswerText']
            
            # Ячейка (2,i): элемент для сопоставления (если есть)
            if row_idx < len(matching_items):
                cell_2_i = table.rows[i].cells[2]
                cell_2_i.text = matching_items[row_idx]['item']
                for paragraph in cell_2_i.paragraphs:
                    paragraph.style = doc.styles['AnswerText']
            else:
                cell_2_i = table.rows[i].cells[2]
                for paragraph in cell_2_i.paragraphs:
                    paragraph.style = doc.styles['AnswerText']
            
            # Ячейка (3,i): правильный ответ (если есть)
            if row_idx < len(matching_items):
                cell_3_i = table.rows[i].cells[3]
                cell_3_i.text = matching_items[row_idx]['answer']
                for paragraph in cell_3_i.paragraphs:
                    paragraph.style = doc.styles['AnswerText']
            else:
                cell_3_i = table.rows[i].cells[3]
                for paragraph in cell_3_i.paragraphs:
                    paragraph.style = doc.styles['AnswerText']

        try:
            cols = ((cfg or {}).get("layout") or {}).get("matching", {}).get("table_cols_pct")
            if isinstance(cols, list):
                QuestionTemplate.apply_table_widths(doc, table, [int(x) for x in cols])
        except Exception:
            pass
    
    @staticmethod
    def render(doc: Document, question: Question, metadata: DocumentMetadata, task_number: int, cfg: Optional[dict] = None):
        """
        Рендеринг вопроса типа matching
        
        Args:
            doc: Объект документа Word
            question: Объект вопроса
            metadata: Метаданные документа
            task_number: Номер задания
        """
        # Заголовок задания
        header_text = (
            f"- Задание {task_number} "
            f"({metadata.pk_prefix}-{metadata.pk_id} – "
            f"{metadata.ipk_prefix}-{metadata.ipk_id} "
            f"{metadata.description})"
        )
        
        doc.add_paragraph(header_text, style='QuestionHeader')
        
        # Текст вопроса
        doc.add_paragraph(question.question_text, style='Question')
        
        # Таблица для ответов matching
        MatchingTemplate.add_matching_answer_table(
            doc,
            task_number,
            question.matching_items,
            question.matching_answers,
            cfg,
        )


class TruefalseTemplate(QuestionTemplate):
    """Шаблон для вопросов типа truefalse"""
    
    @staticmethod
    def add_truefalse_answer_table(doc: Document, task_number: int, answers: List[str], correct_answers: List[str], cfg: Optional[dict] = None):
        """
        Добавление таблицы для ответов truefalse с границами
        
        Args:
            doc: Объект документа Word
            task_number: Номер задания
            answers: Список всех вариантов ответов (Верно, Неверно)
            correct_answers: Список правильных ответов
        """
        # Количество строк: 1 заголовочная + количество вариантов ответов (обычно 2: Верно и Неверно)
        num_rows = 1 + len(answers) if answers else 1
        table = doc.add_table(rows=num_rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        
        # Установка границ таблицы
        QuestionTemplate.set_table_borders(table)
        
        # Заполнение таблицы
        # Ячейка (0,0): номер вопроса
        cell_0_0 = table.rows[0].cells[0]
        cell_0_0.text = f"№{task_number}"
        for paragraph in cell_0_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
        
        # Ячейка (1,0): "Варианты ответа:" (жирный)
        cell_1_0 = table.rows[0].cells[1]
        cell_1_0.text = "Варианты ответа:"
        for paragraph in cell_1_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
            for run in paragraph.runs:
                run.bold = True
        
        # Ячейка (2,0): "Правильный ответ:" (жирный)
        cell_2_0 = table.rows[0].cells[2]
        cell_2_0.text = "Правильный ответ:"
        for paragraph in cell_2_0.paragraphs:
            paragraph.style = doc.styles['AnswerText']
            for run in paragraph.runs:
                run.bold = True
        
        # Заполнение вариантов ответов и правильных ответов
        for i, answer in enumerate(answers, start=1):
            # Ячейка (1,i): вариант ответа
            cell_1_i = table.rows[i].cells[1]
            cell_1_i.text = answer
            for paragraph in cell_1_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']
            
            # Ячейка (2,i): правильный ответ (если этот вариант правильный)
            cell_2_i = table.rows[i].cells[2]
            if answer in correct_answers:
                cell_2_i.text = answer
            else:
                cell_2_i.text = ""  # Пустая строка для неправильных ответов
            for paragraph in cell_2_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']
        
        # Ячейка (0,i): пустая для всех строк с ответами
        for i in range(1, num_rows):
            cell_0_i = table.rows[i].cells[0]
            for paragraph in cell_0_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']

        try:
            # если не задано — используем настройки multichoice
            layout = (cfg or {}).get("layout") or {}
            cols = (layout.get("truefalse") or layout.get("multichoice") or {}).get("table_cols_pct")
            if isinstance(cols, list):
                QuestionTemplate.apply_table_widths(doc, table, [int(x) for x in cols])
        except Exception:
            pass
    
    @staticmethod
    def render(doc: Document, question: Question, metadata: DocumentMetadata, task_number: int, cfg: Optional[dict] = None):
        """
        Рендеринг вопроса типа truefalse
        
        Args:
            doc: Объект документа Word
            question: Объект вопроса
            metadata: Метаданные документа
            task_number: Номер задания
        """
        # Заголовок задания
        header_text = (
            f"- Задание {task_number} "
            f"({metadata.pk_prefix}-{metadata.pk_id} – "
            f"{metadata.ipk_prefix}-{metadata.ipk_id} "
            f"{metadata.description})"
        )
        
        doc.add_paragraph(header_text, style='QuestionHeader')
        
        # Текст вопроса
        doc.add_paragraph(question.question_text, style='Question')
        
        # Таблица для ответов truefalse
        TruefalseTemplate.add_truefalse_answer_table(
            doc,
            task_number,
            question.answers,
            question.correct_answers,
            cfg,
        )


class ShortanswerTemplate(QuestionTemplate):
    """Шаблон для вопросов типа shortanswer"""
    
    @staticmethod
    def add_shortanswer_answer_table(doc: Document, task_number: int, answers: List[str], cfg: Optional[dict] = None):
        """
        Добавление таблицы для ответов shortanswer с границами
        Каждый ответ выводится на отдельной строке
        
        Args:
            doc: Объект документа Word
            task_number: Номер задания
            answers: Список правильных ответов
        """
        # Если нет ответов, создаем пустую таблицу
        if not answers:
            num_rows = 1
        else:
            num_rows = len(answers)
        
        table = doc.add_table(rows=num_rows, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        
        # Установка границ таблицы
        QuestionTemplate.set_table_borders(table)
        
        # Заполнение строк с ответами
        for i, answer in enumerate(answers):
            # Ячейка (0,i): номер вопроса (только в первой строке)
            cell_0_i = table.rows[i].cells[0]
            if i == 0:
                cell_0_i.text = f"№{task_number}"
            # Для остальных строк ячейка остается пустой
            for paragraph in cell_0_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']
            
            # Ячейка (1,i): пустая
            cell_1_i = table.rows[i].cells[1]
            for paragraph in cell_1_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']
            
            # Ячейка (2,i): ответ
            cell_2_i = table.rows[i].cells[2]
            cell_2_i.text = answer
            for paragraph in cell_2_i.paragraphs:
                paragraph.style = doc.styles['AnswerText']

        try:
            cols = ((cfg or {}).get("layout") or {}).get("shortanswer", {}).get("table_cols_pct")
            if isinstance(cols, list):
                QuestionTemplate.apply_table_widths(doc, table, [int(x) for x in cols])
        except Exception:
            pass
    
    @staticmethod
    def render(doc: Document, question: Question, metadata: DocumentMetadata, task_number: int, cfg: Optional[dict] = None):
        """
        Рендеринг вопроса типа shortanswer
        
        Args:
            doc: Объект документа Word
            question: Объект вопроса
            metadata: Метаданные документа
            task_number: Номер задания
        """
        # Заголовок задания
        header_text = (
            f"- Задание {task_number} "
            f"({metadata.pk_prefix}-{metadata.pk_id} – "
            f"{metadata.ipk_prefix}-{metadata.ipk_id} "
            f"{metadata.description})"
        )
        
        doc.add_paragraph(header_text, style='QuestionHeader')
        
        # Текст вопроса
        doc.add_paragraph(question.question_text, style='Question')
        
        # Таблица для ответов shortanswer (каждый ответ на отдельной строке)
        ShortanswerTemplate.add_shortanswer_answer_table(
            doc,
            task_number,
            question.correct_answers if question.correct_answers else [question.reference_answer] if question.reference_answer else [],
            cfg,
        )


class TemplateFactory:
    """Фабрика для создания шаблонов вопросов"""
    
    _templates = {
        'essay_gigachat': EssayGigachatTemplate,
        'shortanswer': ShortanswerTemplate,
        'multichoice': MultichoiceTemplate,
        'matching': MatchingTemplate,
        'truefalse': TruefalseTemplate,
        # Здесь можно добавить другие типы шаблонов
    }
    
    @classmethod
    def get_template(cls, question_type: str):
        """
        Получение шаблона по типу вопроса
        
        Args:
            question_type: Тип вопроса
            
        Returns:
            Класс шаблона или None
        """
        return cls._templates.get(question_type)
    
    @classmethod
    def register_template(cls, question_type: str, template_class):
        """
        Регистрация нового шаблона
        
        Args:
            question_type: Тип вопроса
            template_class: Класс шаблона
        """
        cls._templates[question_type] = template_class

